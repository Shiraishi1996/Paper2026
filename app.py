import csv
import hashlib
import io
import json
import math
import os
import re
import tempfile
import textwrap
import zipfile
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from xml.etree import ElementTree as ET

import matplotlib.pyplot as plt
import pandas as pd
import requests
import streamlit as st
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageOps
from pypdf import PdfReader

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_OK = True
except Exception:
    SKLEARN_OK = False

try:
    import openpyxl  # noqa: F401
    OPENPYXL_OK = True
except Exception:
    OPENPYXL_OK = False

try:
    import pytesseract
    from pytesseract import Output as TesseractOutput
    TESSERACT_OK = True
except Exception:
    TESSERACT_OK = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_OK = True
except Exception:
    PDF2IMAGE_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
    REPORTLAB_OK = True
except Exception:
    REPORTLAB_OK = False


st.set_page_config(
    page_title="PaperForge Deterministic Web",
    layout="wide",
    page_icon="🔬",
    initial_sidebar_state="collapsed",
)


# -----------------------------
# Data models
# -----------------------------
@dataclass
class Evidence:
    source_file: str
    locator: str
    text: str
    kind: str = "text"


@dataclass
class FigureAsset:
    source_file: str
    path: str
    original_name: str
    context: str = ""
    ocr_text: str = ""
    width: int = 0
    height: int = 0
    assigned_section: str = "Figures / Appendix"
    similarity: float = 0.0


@dataclass
class ParsedMaterial:
    name: str
    kind: str
    sha256: str
    raw_text: str = ""
    sections: Dict[str, List[Evidence]] = field(default_factory=dict)
    evidences: List[Evidence] = field(default_factory=list)
    figures: List[FigureAsset] = field(default_factory=list)
    tables: List[pd.DataFrame] = field(default_factory=list)
    table_labels: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


# -----------------------------
# Constants / deterministic rules
# -----------------------------
CANONICAL_SECTIONS = [
    "Title",
    "Abstract",
    "Keywords",
    "Introduction",
    "Methods",
    "Results",
    "Discussion",
    "Conclusion",
    "References",
]

SECTION_PATTERNS = {
    "Abstract": [r"^abstract$", r"^summary$", r"^概要$", r"^要旨$", r"^抄録$"],
    "Keywords": [r"^keywords?$", r"^key\s*words?$", r"^キーワード$"],
    "Introduction": [r"^introduction$", r"^background$", r"^はじめに$", r"^序論$", r"^緒言$"],
    "Methods": [
        r"^materials?\s*(and|&)\s*methods?$", r"^methods?$", r"^methodology$",
        r"^data\s*(and|&)\s*methods?$", r"^手法$", r"^方法$", r"^研究方法$", r"^解析方法$", r"^データと方法$"
    ],
    "Results": [r"^results?$", r"^結果$", r"^解析結果$", r"^実験結果$"],
    "Discussion": [r"^discussion$", r"^考察$"],
    "Conclusion": [r"^conclusions?$", r"^concluding\s*remarks$", r"^結論$", r"^まとめ$"],
    "References": [r"^references?$", r"^bibliography$", r"^参考文献$", r"^引用文献$"],
}

HEADING_PREFIX = re.compile(r"^\s*(?:\d+(?:\.\d+)*[\s.)、-]*)?")
DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", re.I)

# A deliberately small, transparent catalog. Scores are topical similarity only.
# No acceptance probability is produced.
BUILTIN_JOURNALS = [
    {
        "journal": "Natural Hazards Research",
        "scope": "natural hazards disaster risk resilience earthquake flood landslide disaster assessment remote sensing GIS emergency response",
        "publisher": "Elsevier / KeAi",
    },
    {
        "journal": "International Journal of Disaster Risk Reduction",
        "scope": "disaster risk reduction resilience emergency management hazards vulnerability response recovery policy technology",
        "publisher": "Elsevier",
    },
    {
        "journal": "Remote Sensing",
        "scope": "remote sensing satellite UAV image processing GIS earth observation disaster monitoring mapping",
        "publisher": "MDPI",
    },
    {
        "journal": "ISPRS International Journal of Geo-Information",
        "scope": "GIS geoinformation spatial analysis mapping geospatial data web GIS remote sensing spatial decision support",
        "publisher": "MDPI",
    },
    {
        "journal": "IEEE Access",
        "scope": "engineering computer science artificial intelligence image processing robotics control sensing applied technology",
        "publisher": "IEEE",
    },
    {
        "journal": "Regional Studies in Marine Science",
        "scope": "marine science fisheries aquaculture ocean coastal engineering underwater systems environmental monitoring",
        "publisher": "Elsevier",
    },
    {
        "journal": "Finite Elements in Analysis and Design",
        "scope": "finite element methods numerical mechanics structural analysis computational mechanics optimization engineering design",
        "publisher": "Elsevier",
    },
    {
        "journal": "International Journal for Computational Methods in Engineering Science and Mechanics",
        "scope": "computational methods engineering mechanics finite elements numerical analysis optimization simulation",
        "publisher": "Taylor & Francis",
    },
]


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def normalize_line(s: str) -> str:
    return re.sub(r"\s+", " ", s.replace("\u3000", " ")).strip()


def heading_to_canonical(line: str) -> Optional[str]:
    s = normalize_line(line)
    s = HEADING_PREFIX.sub("", s)
    s = s.strip(" :：.-–—")
    low = s.lower()
    for canonical, patterns in SECTION_PATTERNS.items():
        for p in patterns:
            if re.match(p, low, re.I):
                return canonical
    return None


def split_text_into_sections(text: str, source_file: str, locator_prefix: str = "line") -> Tuple[Dict[str, List[Evidence]], List[Evidence]]:
    sections: Dict[str, List[Evidence]] = {k: [] for k in CANONICAL_SECTIONS}
    all_evidence: List[Evidence] = []
    current = "Introduction"  # conservative fallback bucket, not a generated claim
    buffer: List[Tuple[int, str]] = []

    def flush():
        nonlocal buffer
        if not buffer:
            return
        chunk = "\n".join(x[1] for x in buffer).strip()
        if chunk:
            start, end = buffer[0][0], buffer[-1][0]
            ev = Evidence(source_file, f"{locator_prefix} {start}-{end}", chunk)
            sections.setdefault(current, []).append(ev)
            all_evidence.append(ev)
        buffer = []

    for idx, raw in enumerate(text.splitlines(), start=1):
        line = raw.rstrip()
        canonical = heading_to_canonical(line)
        if canonical:
            flush()
            current = canonical
            continue
        if line.strip():
            buffer.append((idx, line))
        elif buffer:
            flush()
    flush()
    return sections, all_evidence


def guess_title_from_text(text: str) -> str:
    for line in text.splitlines()[:40]:
        s = normalize_line(line)
        if 8 <= len(s) <= 180 and not heading_to_canonical(s):
            if not re.match(r"^(doi:|https?://|www\.)", s, re.I):
                return s
    return "（タイトル未抽出）"


# -----------------------------
# PDF
# -----------------------------
def _ocr_best(im: Image.Image) -> Tuple[str, float]:
    """Deterministic OCR: compare PSM 6/11 and keep higher mean confidence."""
    if not TESSERACT_OK:
        return "", -1.0
    try:
        langs = set(pytesseract.get_languages(config=""))
        lang = "jpn+eng" if "jpn" in langs and "eng" in langs else ("jpn" if "jpn" in langs else "eng")
    except Exception:
        lang = "eng"

    img = ImageOps.autocontrast(ImageOps.grayscale(im))
    w, h = img.size
    if max(w, h) < 2200:
        scale = min(3.0, 2200.0 / max(w, h))
        img = img.resize((max(1, int(w * scale)), max(1, int(h * scale))))

    best_text, best_score = "", -1.0
    for psm in (6, 11):
        try:
            data = pytesseract.image_to_data(
                img, lang=lang, config=f"--oem 1 --psm {psm}", output_type=TesseractOutput.DICT
            )
            texts, confs = [], []
            for t, c in zip(data.get("text", []), data.get("conf", [])):
                t = str(t).strip()
                try:
                    cf = float(c)
                except Exception:
                    cf = -1.0
                if t:
                    texts.append(t)
                    if cf >= 0:
                        confs.append(cf)
            text = " ".join(texts).strip()
            score = sum(confs) / len(confs) if confs else -1.0
            # Confidence first, then amount of recovered text as deterministic tie-breaker.
            if score > best_score or (abs(score - best_score) < 1e-9 and len(text) > len(best_text)):
                best_text, best_score = text, score
        except Exception:
            continue
    return best_text, best_score


def parse_pdf(name: str, data: bytes, use_ocr: bool = False) -> ParsedMaterial:
    m = ParsedMaterial(name=name, kind="pdf", sha256=sha256_bytes(data))
    try:
        reader = PdfReader(io.BytesIO(data))
        page_texts: List[str] = []
        for page in reader.pages:
            page_texts.append((page.extract_text() or "").strip())

        if use_ocr and TESSERACT_OK and PDF2IMAGE_OK:
            for i, txt in enumerate(list(page_texts), start=1):
                if len(normalize_line(txt)) >= 80:
                    continue
                try:
                    imgs = convert_from_bytes(
                        data, dpi=300, first_page=i, last_page=i, fmt="png", thread_count=1
                    )
                    if imgs:
                        ocr_text, conf = _ocr_best(imgs[0])
                        if len(normalize_line(ocr_text)) > len(normalize_line(txt)):
                            page_texts[i - 1] = ocr_text
                            m.evidences.append(Evidence(name, f"page {i} OCR confidence={conf:.1f}", ocr_text, kind="pdf_ocr"))
                except Exception as e:
                    m.warnings.append(f"PDF page {i} OCR失敗: {e}")

        pages = []
        for i, txt in enumerate(page_texts, start=1):
            if txt:
                pages.append(f"\n[PAGE {i}]\n{txt}")
                m.evidences.append(Evidence(name, f"page {i}", txt))
        m.raw_text = "\n".join(pages)
        if not m.raw_text.strip():
            m.warnings.append("PDFからテキストを抽出できませんでした。OCRを有効にしてください。")
        m.sections, ev = split_text_into_sections(m.raw_text, name, "PDF text line")
        m.evidences.extend(ev)
    except Exception as e:
        m.warnings.append(f"PDF解析エラー: {e}")
    return m


# -----------------------------
# DOCX: paragraphs, tables, embedded images + local context
# -----------------------------
def _docx_image_context(data: bytes, temp_dir: str, source_name: str) -> List[FigureAsset]:
    out: List[FigureAsset] = []
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            if "word/document.xml" not in z.namelist():
                return out
            relmap = {}
            if "word/_rels/document.xml.rels" in z.namelist():
                relroot = ET.fromstring(z.read("word/_rels/document.xml.rels"))
                for rel in relroot:
                    rid = rel.attrib.get("Id")
                    target = rel.attrib.get("Target", "")
                    if rid and "media/" in target:
                        relmap[rid] = "word/" + target.lstrip("/")

            root = ET.fromstring(z.read("word/document.xml"))
            body = root.find("w:body", ns)
            if body is None:
                return out

            blocks = []
            for child in list(body):
                tag = child.tag.split("}")[-1]
                if tag != "p":
                    continue
                texts = [t.text or "" for t in child.findall(".//w:t", ns)]
                paragraph_text = "".join(texts).strip()
                rids = [b.attrib.get(f"{{{ns['r']}}}embed") for b in child.findall(".//a:blip", ns)]
                blocks.append((paragraph_text, [x for x in rids if x]))

            for i, (txt, rids) in enumerate(blocks):
                if not rids:
                    continue
                prev_text = blocks[i - 1][0] if i > 0 else ""
                next_text = blocks[i + 1][0] if i + 1 < len(blocks) else ""
                context = " | ".join(x for x in [prev_text, txt, next_text] if x)
                for rid in rids:
                    member = relmap.get(rid)
                    if not member or member not in z.namelist():
                        continue
                    raw = z.read(member)
                    ext = Path(member).suffix or ".bin"
                    fname = f"{Path(source_name).stem}_{rid}{ext}"
                    dest = os.path.join(temp_dir, fname)
                    with open(dest, "wb") as f:
                        f.write(raw)
                    width = height = 0
                    try:
                        with Image.open(io.BytesIO(raw)) as im:
                            width, height = im.size
                    except Exception:
                        pass
                    out.append(FigureAsset(source_name, dest, fname, context=context, width=width, height=height))
    except Exception:
        pass
    return out


def parse_docx(name: str, data: bytes, temp_dir: str) -> ParsedMaterial:
    m = ParsedMaterial(name=name, kind="docx", sha256=sha256_bytes(data))
    try:
        doc = Document(io.BytesIO(data))
        paras = []
        for i, p in enumerate(doc.paragraphs, start=1):
            t = p.text.strip()
            if t:
                paras.append(t)
                m.evidences.append(Evidence(name, f"paragraph {i}", t))
        m.raw_text = "\n".join(paras)
        m.sections, sec_ev = split_text_into_sections(m.raw_text, name, "DOCX text line")
        m.evidences.extend(sec_ev)

        for ti, table in enumerate(doc.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                max_len = max(len(r) for r in rows)
                rows = [r + [""] * (max_len - len(r)) for r in rows]
                header = rows[0]
                if len(set(header)) != len(header) or any(not h for h in header):
                    header = [f"col_{i+1}" for i in range(max_len)]
                    body = rows
                else:
                    body = rows[1:]
                df = pd.DataFrame(body, columns=header)
                m.tables.append(df)
                m.table_labels.append(f"{name} / table {ti}")
                m.evidences.append(Evidence(name, f"table {ti}", df.to_csv(index=False), kind="table"))

        m.figures.extend(_docx_image_context(data, temp_dir, name))
    except Exception as e:
        m.warnings.append(f"DOCX解析エラー: {e}")
    return m


# -----------------------------
# CSV / TSV / XLSX
# -----------------------------
def _read_csv_bytes(data: bytes, name: str) -> pd.DataFrame:
    encodings = ["utf-8-sig", "utf-8", "cp932", "shift_jis", "latin1"]
    last = None
    for enc in encodings:
        try:
            text = data.decode(enc)
            sample = text[:4096]
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
                sep = dialect.delimiter
            except Exception:
                sep = "\t" if name.lower().endswith(".tsv") else ","
            return pd.read_csv(io.StringIO(text), sep=sep)
        except Exception as e:
            last = e
    raise last or ValueError("CSV decode failed")


def parse_table_file(name: str, data: bytes) -> ParsedMaterial:
    ext = Path(name).suffix.lower()
    m = ParsedMaterial(name=name, kind="table", sha256=sha256_bytes(data))
    try:
        dfs = []
        labels = []
        if ext in [".csv", ".tsv"]:
            df = _read_csv_bytes(data, name)
            dfs = [df]
            labels = [name]
        elif ext in [".xlsx", ".xls"]:
            if not OPENPYXL_OK and ext == ".xlsx":
                raise RuntimeError("openpyxl が必要です")
            book = pd.ExcelFile(io.BytesIO(data))
            for sheet in book.sheet_names:
                dfs.append(pd.read_excel(book, sheet_name=sheet))
                labels.append(f"{name} / {sheet}")
        else:
            raise ValueError("Unsupported table format")

        for df, label in zip(dfs, labels):
            m.tables.append(df)
            m.table_labels.append(label)
            m.evidences.append(Evidence(name, label, df.head(200).to_csv(index=False), kind="table"))
            m.raw_text += f"\n[{label}]\n{df.head(200).to_csv(index=False)}"
    except Exception as e:
        m.warnings.append(f"表データ解析エラー: {e}")
    return m


# -----------------------------
# TXT / MD
# -----------------------------
def parse_text_file(name: str, data: bytes) -> ParsedMaterial:
    m = ParsedMaterial(name=name, kind="text", sha256=sha256_bytes(data))
    for enc in ["utf-8-sig", "utf-8", "cp932", "shift_jis", "latin1"]:
        try:
            m.raw_text = data.decode(enc)
            break
        except Exception:
            continue
    if not m.raw_text:
        m.warnings.append("テキストの文字コードを判定できませんでした。")
    m.sections, m.evidences = split_text_into_sections(m.raw_text, name)
    return m


# -----------------------------
# Image
# -----------------------------
def parse_image(name: str, data: bytes, temp_dir: str, use_ocr: bool) -> ParsedMaterial:
    m = ParsedMaterial(name=name, kind="image", sha256=sha256_bytes(data))
    dest = os.path.join(temp_dir, Path(name).name)
    with open(dest, "wb") as f:
        f.write(data)
    width = height = 0
    ocr_text = ""
    try:
        with Image.open(io.BytesIO(data)) as im:
            width, height = im.size
            if use_ocr and TESSERACT_OK:
                try:
                    ocr_text, ocr_conf = _ocr_best(im)
                    if ocr_text:
                        m.warnings.append(f"OCR平均信頼度: {ocr_conf:.1f}")
                except Exception as e:
                    m.warnings.append(f"OCR失敗: {e}")
    except Exception as e:
        m.warnings.append(f"画像解析エラー: {e}")
    fig = FigureAsset(name, dest, Path(name).name, context=Path(name).stem.replace("_", " "), ocr_text=ocr_text, width=width, height=height)
    m.figures.append(fig)
    if ocr_text:
        m.evidences.append(Evidence(name, "image OCR", ocr_text, kind="image_ocr"))
    return m


# -----------------------------
# Deterministic statistics
# -----------------------------
def numeric_summary(df: pd.DataFrame) -> pd.DataFrame:
    num = df.select_dtypes(include="number")
    if num.empty:
        return pd.DataFrame()
    desc = num.describe(percentiles=[0.25, 0.5, 0.75]).T
    desc["missing"] = num.isna().sum()
    desc["missing_rate"] = num.isna().mean()
    desc["median"] = num.median(numeric_only=True)
    desc["iqr"] = num.quantile(0.75) - num.quantile(0.25)
    cols = [c for c in ["count", "missing", "missing_rate", "mean", "std", "min", "25%", "median", "75%", "max", "iqr"] if c in desc.columns]
    return desc[cols]


def make_table_figures(materials: List[ParsedMaterial], temp_dir: str) -> List[FigureAsset]:
    figs: List[FigureAsset] = []
    idx = 0
    for m in materials:
        for label, df in zip(m.table_labels, m.tables):
            num = df.select_dtypes(include="number")
            if num.empty:
                continue
            # Distribution plot: max 6 columns to prevent unreadable figures.
            cols = list(num.columns[:6])
            if not cols:
                continue
            idx += 1
            p = os.path.join(temp_dir, f"auto_distribution_{idx}.png")
            fig, ax = plt.subplots(figsize=(8, 4.8), dpi=220)
            num[cols].plot(kind="box", ax=ax)
            ax.set_title(f"Distribution: {label}")
            ax.set_ylabel("Value")
            ax.grid(True, linestyle="--", alpha=0.35)
            plt.tight_layout()
            plt.savefig(p, bbox_inches="tight")
            plt.close(fig)
            figs.append(FigureAsset(m.name, p, Path(p).name, context=f"numeric distribution {label}", assigned_section="Results"))

            if len(cols) >= 2:
                idx += 1
                p2 = os.path.join(temp_dir, f"auto_correlation_{idx}.png")
                corr = num[cols].corr(method="spearman")
                fig2, ax2 = plt.subplots(figsize=(6.5, 5.5), dpi=220)
                im = ax2.imshow(corr.values, vmin=-1, vmax=1, aspect="auto")
                ax2.set_xticks(range(len(cols)), cols, rotation=45, ha="right")
                ax2.set_yticks(range(len(cols)), cols)
                ax2.set_title(f"Spearman correlation: {label}")
                fig2.colorbar(im, ax=ax2, shrink=0.85)
                plt.tight_layout()
                plt.savefig(p2, bbox_inches="tight")
                plt.close(fig2)
                figs.append(FigureAsset(m.name, p2, Path(p2).name, context=f"spearman correlation {label}", assigned_section="Results"))
    return figs


# -----------------------------
# Figure placement: TF-IDF over filename/OCR/context vs section source text
# -----------------------------
def assign_figures(figures: List[FigureAsset], section_texts: Dict[str, str]) -> None:
    target_sections = ["Introduction", "Methods", "Results", "Discussion", "Conclusion"]
    section_docs = [section_texts.get(s, "") for s in target_sections]
    if not SKLEARN_OK or not any(x.strip() for x in section_docs):
        for f in figures:
            if f.assigned_section == "Figures / Appendix":
                f.assigned_section = "Results"
        return

    for fig in figures:
        if fig.assigned_section != "Figures / Appendix":
            continue
        fig_doc = " ".join([fig.original_name, fig.context, fig.ocr_text]).strip()
        if not fig_doc:
            fig.assigned_section = "Figures / Appendix"
            continue
        docs = section_docs + [fig_doc]
        try:
            vec = TfidfVectorizer(analyzer="char_wb", ngram_range=(2, 5), min_df=1).fit_transform(docs)
            sims = cosine_similarity(vec[-1], vec[:-1]).flatten()
            best = int(sims.argmax())
            score = float(sims[best])
            fig.similarity = score
            fig.assigned_section = target_sections[best] if score >= 0.06 else "Figures / Appendix"
        except Exception:
            fig.assigned_section = "Figures / Appendix"


# -----------------------------
# References: DOI extraction + optional Crossref verification
# -----------------------------
def extract_dois(text: str) -> List[str]:
    found = []
    seen = set()
    for m in DOI_RE.finditer(text or ""):
        doi = m.group(0).rstrip(".,;)]}").lower()
        if doi not in seen:
            seen.add(doi)
            found.append(doi)
    return found


def crossref_lookup(doi: str) -> Dict[str, str]:
    url = f"https://api.crossref.org/works/{requests.utils.quote(doi, safe='')}"
    r = requests.get(url, timeout=12, headers={"User-Agent": "PaperForgeDeterministic/2.0 (mailto:research@example.invalid)"})
    r.raise_for_status()
    msg = r.json().get("message", {})
    authors = []
    for a in msg.get("author", [])[:20]:
        given = a.get("given", "")
        family = a.get("family", "")
        authors.append((given + " " + family).strip())
    title = (msg.get("title") or [""])[0]
    journal = (msg.get("container-title") or [""])[0]
    year = ""
    for key in ["published-print", "published-online", "issued"]:
        parts = msg.get(key, {}).get("date-parts", [])
        if parts and parts[0]:
            year = str(parts[0][0])
            break
    return {
        "doi": doi,
        "title": title,
        "authors": ", ".join(authors),
        "journal": journal,
        "year": year,
        "url": msg.get("URL", f"https://doi.org/{doi}"),
        "verified": "yes",
    }


def verify_dois(dois: List[str], enabled: bool) -> List[Dict[str, str]]:
    refs = []
    for doi in dois:
        if not enabled:
            refs.append({"doi": doi, "title": "", "authors": "", "journal": "", "year": "", "url": f"https://doi.org/{doi}", "verified": "not checked"})
            continue
        try:
            refs.append(crossref_lookup(doi))
        except Exception as e:
            refs.append({"doi": doi, "title": "", "authors": "", "journal": "", "year": "", "url": f"https://doi.org/{doi}", "verified": f"failed: {e}"})
    return refs


def format_reference(ref: Dict[str, str], idx: int) -> str:
    if ref.get("title"):
        bits = [ref.get("authors", ""), f"({ref.get('year', '')})" if ref.get("year") else "", ref.get("title", ""), ref.get("journal", ""), f"https://doi.org/{ref.get('doi', '')}"]
        return f"[{idx}] " + ". ".join([b for b in bits if b]).replace("..", ".")
    return f"[{idx}] DOI: {ref.get('doi', '')} ({ref.get('verified', '')})"


# -----------------------------
# Corpus / source-constrained assembly
# -----------------------------
def merge_sections(materials: List[ParsedMaterial]) -> Dict[str, List[Evidence]]:
    merged = {k: [] for k in CANONICAL_SECTIONS}
    for m in materials:
        for k, items in m.sections.items():
            merged.setdefault(k, []).extend(items)
    return merged


def section_plain_text(items: List[Evidence]) -> str:
    return "\n\n".join(ev.text.strip() for ev in items if ev.text.strip())


def source_tag(ev: Evidence) -> str:
    return f"[{ev.source_file} | {ev.locator}]"


def source_constrained_section(items: List[Evidence]) -> str:
    if not items:
        return "（素材から該当セクションを抽出できませんでした。生成補完は行っていません。）"
    parts = []
    for ev in items:
        parts.append(f"{source_tag(ev)}\n{ev.text.strip()}")
    return "\n\n".join(parts)


# -----------------------------
# Journal matching: no prediction, only text similarity
# -----------------------------
def load_journal_catalog(uploaded) -> List[Dict[str, str]]:
    if uploaded is None:
        return BUILTIN_JOURNALS
    data = uploaded.getvalue()
    try:
        df = _read_csv_bytes(data, uploaded.name)
        needed = {"journal", "scope"}
        if not needed.issubset(set(df.columns)):
            raise ValueError("CSVには journal, scope 列が必要です")
        out = []
        for _, r in df.iterrows():
            out.append({
                "journal": str(r.get("journal", "")),
                "scope": str(r.get("scope", "")),
                "publisher": str(r.get("publisher", "")),
            })
        return out
    except Exception as e:
        st.warning(f"投稿先カタログの読込に失敗したため内蔵カタログを使用します: {e}")
        return BUILTIN_JOURNALS


def journal_similarity(corpus: str, catalog: List[Dict[str, str]]) -> pd.DataFrame:
    if not corpus.strip() or not catalog or not SKLEARN_OK:
        return pd.DataFrame(columns=["journal", "publisher", "scope_similarity"])
    docs = [corpus] + [x.get("scope", "") for x in catalog]
    vec = TfidfVectorizer(analyzer="char_wb", ngram_range=(2, 5), min_df=1).fit_transform(docs)
    sims = cosine_similarity(vec[0], vec[1:]).flatten()
    rows = []
    for c, s in zip(catalog, sims):
        rows.append({"journal": c.get("journal", ""), "publisher": c.get("publisher", ""), "scope_similarity": float(s)})
    return pd.DataFrame(rows).sort_values("scope_similarity", ascending=False).reset_index(drop=True)


# -----------------------------
# Exports
# -----------------------------
def export_docx(title: str, author: str, affiliation: str, merged: Dict[str, List[Evidence]], figures: List[FigureAsset], materials: List[ParsedMaterial], refs: List[Dict[str, str]]) -> bytes:
    doc = Document()
    doc.add_heading(title or "（タイトル未設定）", level=0)
    if author or affiliation:
        doc.add_paragraph(" / ".join([x for x in [author, affiliation] if x]))

    for section in ["Abstract", "Keywords", "Introduction", "Methods", "Results", "Discussion", "Conclusion"]:
        doc.add_heading(section, level=1)
        items = merged.get(section, [])
        if not items:
            doc.add_paragraph("素材から該当内容を抽出できませんでした。生成補完は行っていません。")
        else:
            for ev in items:
                p = doc.add_paragraph()
                p.add_run(source_tag(ev) + "\n").bold = True
                p.add_run(ev.text)

        for fig in [f for f in figures if f.assigned_section == section]:
            try:
                doc.add_picture(fig.path, width=Inches(5.8))
                cap = doc.add_paragraph()
                cap.alignment = 1
                cap.add_run(f"Figure: {fig.original_name} / source={fig.source_file} / deterministic placement score={fig.similarity:.3f}").italic = True
                if fig.context:
                    doc.add_paragraph(f"Source context: {fig.context}")
            except Exception as e:
                doc.add_paragraph(f"[Figure unavailable: {fig.original_name}: {e}]")

    appendix_figs = [f for f in figures if f.assigned_section == "Figures / Appendix"]
    if appendix_figs:
        doc.add_heading("Figures / Appendix", level=1)
        for fig in appendix_figs:
            try:
                doc.add_picture(fig.path, width=Inches(5.8))
                doc.add_paragraph(f"Figure: {fig.original_name} / source={fig.source_file}")
            except Exception:
                pass

    # Numeric data appendices
    for m in materials:
        for label, df in zip(m.table_labels, m.tables):
            doc.add_heading(f"Data table: {label}", level=2)
            summary = numeric_summary(df)
            if not summary.empty:
                table = doc.add_table(rows=len(summary) + 1, cols=len(summary.columns) + 1)
                table.style = "Table Grid"
                table.cell(0, 0).text = "variable"
                for j, col in enumerate(summary.columns, start=1):
                    table.cell(0, j).text = str(col)
                for i, (idx, row) in enumerate(summary.iterrows(), start=1):
                    table.cell(i, 0).text = str(idx)
                    for j, val in enumerate(row, start=1):
                        try:
                            table.cell(i, j).text = f"{float(val):.6g}"
                        except Exception:
                            table.cell(i, j).text = str(val)

    doc.add_heading("References verified from DOI metadata", level=1)
    if refs:
        for i, ref in enumerate(refs, start=1):
            doc.add_paragraph(format_reference(ref, i))
    else:
        doc.add_paragraph("DOIを素材から抽出できませんでした。参考文献の新規生成は行っていません。")

    doc.add_heading("Provenance / Audit Trail", level=1)
    for m in materials:
        doc.add_paragraph(f"{m.name} | type={m.kind} | SHA-256={m.sha256}")
        for w in m.warnings:
            doc.add_paragraph(f"WARNING: {w}")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def export_markdown(title: str, author: str, affiliation: str, merged: Dict[str, List[Evidence]], figures: List[FigureAsset], refs: List[Dict[str, str]]) -> bytes:
    lines = [f"# {title or '（タイトル未設定）'}", ""]
    if author or affiliation:
        lines += [f"**Author/Affiliation:** {' / '.join([x for x in [author, affiliation] if x])}", ""]
    for section in ["Abstract", "Keywords", "Introduction", "Methods", "Results", "Discussion", "Conclusion"]:
        lines += [f"## {section}", ""]
        lines += [source_constrained_section(merged.get(section, [])), ""]
        for fig in [f for f in figures if f.assigned_section == section]:
            lines += [f"**Figure:** {fig.original_name} (source: {fig.source_file}; score={fig.similarity:.3f})", ""]
    lines += ["## References", ""]
    if refs:
        lines += [format_reference(r, i) for i, r in enumerate(refs, start=1)]
    else:
        lines += ["DOI-based references were not found. No references were fabricated."]
    return "\n".join(lines).encode("utf-8")


def export_latex(title: str, author: str, affiliation: str, merged: Dict[str, List[Evidence]], refs: List[Dict[str, str]]) -> bytes:
    def esc(s: str) -> str:
        rep = {"\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}"}
        for a, b in rep.items():
            s = s.replace(a, b)
        return s

    lines = [
        r"\documentclass[11pt,a4paper]{article}",
        r"\usepackage[margin=25mm]{geometry}",
        r"\usepackage{graphicx}",
        r"\usepackage{booktabs}",
        r"\usepackage{amsmath}",
        rf"\title{{{esc(title or 'Untitled')}}}",
        rf"\author{{{esc(author)}\\{esc(affiliation)}}}",
        r"\date{}",
        r"\begin{document}",
        r"\maketitle",
    ]
    for section in ["Abstract", "Keywords", "Introduction", "Methods", "Results", "Discussion", "Conclusion"]:
        if section == "Abstract":
            lines.append(r"\begin{abstract}")
            lines.append(esc(source_constrained_section(merged.get(section, []))))
            lines.append(r"\end{abstract}")
        else:
            lines.append(rf"\section{{{esc(section)}}}")
            lines.append(esc(source_constrained_section(merged.get(section, []))))
    lines.append(r"\section{References}")
    for i, ref in enumerate(refs, start=1):
        lines.append(esc(format_reference(ref, i)) + r"\\")
    lines.append(r"\end{document}")
    return "\n".join(lines).encode("utf-8")


def export_pdf(title: str, author: str, affiliation: str, merged: Dict[str, List[Evidence]], refs: List[Dict[str, str]]) -> Optional[bytes]:
    if not REPORTLAB_OK:
        return None
    bio = io.BytesIO()
    try:
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("HeiseiMin-W3"))
            font_name = "HeiseiMin-W3"
        except Exception:
            font_name = "Helvetica"
        styles = getSampleStyleSheet()
        body = ParagraphStyle("JPBody", parent=styles["BodyText"], fontName=font_name, fontSize=9.5, leading=14)
        h1 = ParagraphStyle("JPH1", parent=styles["Heading1"], fontName=font_name, fontSize=15, leading=19, spaceAfter=8)
        title_style = ParagraphStyle("JPTitle", parent=styles["Title"], fontName=font_name, fontSize=18, leading=22, alignment=TA_CENTER)
        story = [Paragraph(title or "Untitled", title_style), Spacer(1, 8)]
        if author or affiliation:
            story.append(Paragraph(" / ".join([x for x in [author, affiliation] if x]), body))
            story.append(Spacer(1, 10))
        for section in ["Abstract", "Keywords", "Introduction", "Methods", "Results", "Discussion", "Conclusion"]:
            story.append(Paragraph(section, h1))
            txt = source_constrained_section(merged.get(section, [])).replace("\n", "<br/>")
            story.append(Paragraph(txt, body))
            story.append(Spacer(1, 8))
        story.append(PageBreak())
        story.append(Paragraph("References", h1))
        for i, ref in enumerate(refs, start=1):
            story.append(Paragraph(format_reference(ref, i), body))
        SimpleDocTemplate(bio, pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42).build(story)
        return bio.getvalue()
    except Exception:
        return None


# -----------------------------
# UI
# -----------------------------
st.title("🔬 PaperForge Deterministic Web — LLM 0%")
st.caption(
    "ブラウザ上で動作する決定論的論文構築システム。"
    "生成AI・LLM・OpenAI/Gemini/Claude APIは一切使用せず、入力素材から抽出できた情報だけを構造化・解析・組版します。"
)

with st.expander("精度方針", expanded=False):
    st.markdown(
        """
- **LLM / 生成AI / AI Gateway / Chat API: 使用しません**
- PDF / DOCX / CSV / Excel / TXT / MD / 画像をローカルルールで解析
- DOIは素材から正規表現で抽出し、必要な場合のみCrossrefで実在照合
- 図の配置は、Word内の前後文・画像ファイル名・任意OCRテキストと各セクションのTF-IDF類似度で決定
- 数値はCSV/Excelから直接計算。素材にない数値は作りません
- 投稿先は**採択確率ではなく scope 文面との機械的類似度**だけを表示します
- 各出力に出典ファイル名・ページ/段落番号・SHA-256を残します
        """
    )

left, right = st.columns([1, 1])
with left:
    st.subheader("1. 基本情報")
    title = st.text_input("論文タイトル（既存タイトルを入力。自動生成しません）", value="")
    author = st.text_input("著者名", value="")
    affiliation = st.text_input("所属", value="")

    st.subheader("2. 素材")
    uploads = st.file_uploader(
        "PDF / DOCX / CSV / TSV / XLSX / TXT / MD / PNG / JPG / TIFF / WEBP",
        type=["pdf", "docx", "csv", "tsv", "xlsx", "xls", "txt", "md", "png", "jpg", "jpeg", "tif", "tiff", "webp"],
        accept_multiple_files=True,
    )
    use_ocr = st.checkbox("画像・スキャンPDF OCRを使用（Tesseract、日本語＋英語・LLMではありません）", value=False)
    verify_crossref = st.checkbox(
        "DOIをCrossrefで実在照合（任意・非LLMの書誌API）", value=True
    )

with right:
    st.subheader("3. 投稿先照合")
    journal_catalog_file = st.file_uploader("任意: journal, scope, publisher 列を持つ投稿先CSV", type=["csv"], accept_multiple_files=False)
    st.caption("未指定時は小規模な内蔵カタログを使います。表示値は採択率ではなくスコープ文面の類似度です。")

    st.subheader("4. 出力")
    output_types = st.multiselect("出力形式", ["Word (.docx)", "PDF (.pdf)", "LaTeX (.tex)", "Markdown (.md)"], default=["Word (.docx)"])

build = st.button("⚡ LLMなしで決定論的に構築", type="primary", use_container_width=True)

if build:
    if not uploads:
        st.error("素材ファイルを1つ以上追加してください。")
        st.stop()

    with st.spinner("決定論的に解析中…"):
        with tempfile.TemporaryDirectory() as temp_dir:
            materials: List[ParsedMaterial] = []
            for up in uploads:
                data = up.getvalue()
                ext = Path(up.name).suffix.lower()
                if ext == ".pdf":
                    m = parse_pdf(up.name, data, use_ocr=use_ocr)
                elif ext == ".docx":
                    m = parse_docx(up.name, data, temp_dir)
                elif ext in [".csv", ".tsv", ".xlsx", ".xls"]:
                    m = parse_table_file(up.name, data)
                elif ext in [".txt", ".md"]:
                    m = parse_text_file(up.name, data)
                elif ext in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp"]:
                    m = parse_image(up.name, data, temp_dir, use_ocr)
                else:
                    m = ParsedMaterial(up.name, "unknown", sha256_bytes(data), warnings=["未対応形式"])
                materials.append(m)

            merged = merge_sections(materials)
            corpus = "\n".join(m.raw_text for m in materials if m.raw_text)
            if not title:
                title = guess_title_from_text(corpus)

            figures = []
            for m in materials:
                figures.extend(m.figures)
            figures.extend(make_table_figures(materials, temp_dir))
            section_texts = {s: section_plain_text(merged.get(s, [])) for s in CANONICAL_SECTIONS}
            assign_figures(figures, section_texts)

            dois = extract_dois(corpus)
            refs = verify_dois(dois, verify_crossref)

            catalog = load_journal_catalog(journal_catalog_file)
            journal_df = journal_similarity(corpus, catalog)

            st.success("解析完了。生成文章は作成していません。")

            tabs = st.tabs(["構造化本文", "数値解析", "図", "参考文献", "投稿先", "監査ログ"])

            with tabs[0]:
                st.markdown(f"# {title}")
                for sec in ["Abstract", "Keywords", "Introduction", "Methods", "Results", "Discussion", "Conclusion"]:
                    st.markdown(f"## {sec}")
                    items = merged.get(sec, [])
                    if not items:
                        st.info("素材から該当内容を抽出できませんでした。生成補完は行っていません。")
                    for ev in items:
                        st.caption(source_tag(ev))
                        st.text(ev.text)

            with tabs[1]:
                any_table = False
                for m in materials:
                    for label, df in zip(m.table_labels, m.tables):
                        any_table = True
                        st.markdown(f"### {label}")
                        st.dataframe(df.head(500), use_container_width=True)
                        summ = numeric_summary(df)
                        if not summ.empty:
                            st.markdown("**記述統計**")
                            st.dataframe(summ, use_container_width=True)
                            if df.select_dtypes(include="number").shape[1] >= 2:
                                st.markdown("**Spearman相関**")
                                st.dataframe(df.select_dtypes(include="number").corr(method="spearman"), use_container_width=True)
                if not any_table:
                    st.info("数値表データはありません。")

            with tabs[2]:
                if figures:
                    for f in figures:
                        st.markdown(f"**{f.original_name}** → `{f.assigned_section}` (score={f.similarity:.3f})")
                        try:
                            st.image(f.path, caption=f"source={f.source_file} | context={f.context[:200]}", use_container_width=True)
                        except Exception:
                            st.warning("画像プレビュー不可")
                else:
                    st.info("画像はありません。")

            with tabs[3]:
                if refs:
                    ref_df = pd.DataFrame(refs)
                    st.dataframe(ref_df, use_container_width=True)
                else:
                    st.info("素材からDOIを抽出できませんでした。参考文献は新規生成しません。")

            with tabs[4]:
                if journal_df.empty:
                    st.info("投稿先類似度を計算できませんでした。")
                else:
                    shown = journal_df.head(10).copy()
                    shown["scope_similarity"] = shown["scope_similarity"].map(lambda x: round(x, 4))
                    st.dataframe(shown, use_container_width=True)
                    st.caption("注意: この値は文面類似度であり、採択確率・査読通過率・最新の投稿要件を意味しません。")

            with tabs[5]:
                for m in materials:
                    st.code(f"{m.name}\ntype={m.kind}\nsha256={m.sha256}")
                    for w in m.warnings:
                        st.warning(f"{m.name}: {w}")

            st.divider()
            st.subheader("出力")
            if "Word (.docx)" in output_types:
                word_bytes = export_docx(title, author, affiliation, merged, figures, materials, refs)
                st.download_button("📥 Word (.docx)", word_bytes, file_name="paperforge_deterministic.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

            if "Markdown (.md)" in output_types:
                md_bytes = export_markdown(title, author, affiliation, merged, figures, refs)
                st.download_button("📥 Markdown (.md)", md_bytes, file_name="paperforge_deterministic.md", mime="text/markdown", use_container_width=True)

            if "LaTeX (.tex)" in output_types:
                tex_bytes = export_latex(title, author, affiliation, merged, refs)
                st.download_button("📥 LaTeX (.tex)", tex_bytes, file_name="paperforge_deterministic.tex", mime="text/x-tex", use_container_width=True)

            if "PDF (.pdf)" in output_types:
                pdf_bytes = export_pdf(title, author, affiliation, merged, refs)
                if pdf_bytes:
                    st.download_button("📥 PDF (.pdf)", pdf_bytes, file_name="paperforge_deterministic.pdf", mime="application/pdf", use_container_width=True)
                else:
                    st.warning("PDF出力にはreportlabが必要です。Word/LaTeX出力は利用できます。")
